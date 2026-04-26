from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "report.docx"

UNIVERSITY = "Министерство науки и высшего образования Российской Федерации"
ORGANIZATION = "Федеральное государственное автономное образовательное учреждение высшего образования"
INSTITUTE = "«Московский авиационный институт (национальный исследовательский университет)»"
DEPARTMENT = "Кафедра 806 «Фундаментальная информатика и информационные технологии»"
STUDENT = "Выполнили студенты:"
STUDENT_1 = "Паньков Михаил"
STUDENT_2 = "Губенко Константин"
GROUP = "Группа: М8О-212БВ-24"
TEACHER = "Проверил: ____________________"
CITY_YEAR = "Москва, 2026"


def set_run(run, font="Times New Roman", size=14, bold=False, italic=False):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_doc_defaults(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)


def set_margins(section):
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.different_first_page_header_footer = True


def add_page_number(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=True, size=14):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(0)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    set_run(r, size=size)
    return p


def heading(doc, number, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"{number} {text}" if number else text)
    set_run(r, size=16, bold=True)
    return p


def caption(doc, text, centered=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run(r, size=12, italic=True)


def toc_line(doc, title, page):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(0)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(16.0))
    r = p.add_run(f"{title}\t{page}")
    set_run(r, size=14)


def add_toc(doc):
    for title, page in [
        ("Введение", "3"),
        ("1 Постановка задачи", "3"),
        ("2 Описание архитектуры", "3"),
        ("3 Индексирование B+-tree", "4"),
        ("4 Парсер и выполнение запросов", "4"),
        ("5 Надежность и журналирование", "5"),
        ("6 Тестирование", "5"),
        ("7 Руководство пользователя", "6"),
        ("Вывод", "6"),
        ("Список использованных источников", "6"),
        ("Приложение А", "7"),
    ]:
        toc_line(doc, title, page)


def add_page_break(doc):
    doc.add_page_break()


def spacer(doc, height_pt=12):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("")
    set_run(r, size=max(1, height_pt))


def add_title_page(doc):
    for text in [UNIVERSITY, ORGANIZATION, INSTITUTE, DEPARTMENT]:
        paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)

    for _ in range(3):
        spacer(doc, 10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        "ПОЯСНИТЕЛЬНАЯ ЗАПИСКА",
        "к курсовой работе",
        "по дисциплине «Системное программирование»",
        "на тему: «Система управления базами данных с индексом B+-tree»",
        "Вариант 2",
    ]:
        r = p.add_run(line + "\n")
        set_run(r, size=16, bold=(line == "ПОЯСНИТЕЛЬНАЯ ЗАПИСКА"))

    for _ in range(3):
        spacer(doc, 10)

    for text in [STUDENT, STUDENT_1, STUDENT_2, GROUP, TEACHER]:
        paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.RIGHT, first_indent=False)

    for _ in range(3):
        spacer(doc, 10)

    paragraph(doc, CITY_YEAR, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)


def add_table(doc):
    caption(doc, "Таблица 1. Поддерживаемые команды")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cells = table.rows[0].cells
    for i, text in enumerate(["Группа", "Команды", "Назначение"]):
        cells[i].text = text
    rows = [
        ("Метаданные", "CREATE DATABASE, DROP DATABASE, USE", "Управление базами данных и текущим контекстом"),
        ("DDL", "CREATE TABLE, DROP TABLE", "Создание и удаление схем таблиц"),
        ("DML", "INSERT, UPDATE, DELETE, SELECT", "Манипулирование строками и выборка данных"),
        ("История", "REVERT", "Откат состояния таблицы к моменту времени"),
    ]
    for row_data in rows:
        row = table.add_row().cells
        for i, text in enumerate(row_data):
            row[i].text = text
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    set_run(run, size=14)


def add_listing(doc, title, text):
    caption(doc, title)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.keep_together = True
    run = p.add_run(text)
    set_run(run, font="Consolas", size=12)


def build():
    doc = Document()
    set_doc_defaults(doc)
    set_margins(doc.sections[0])
    add_page_number(doc.sections[0])

    add_title_page(doc)
    add_page_break(doc)

    heading(doc, "", "Содержание")
    add_toc(doc)
    add_page_break(doc)

    heading(doc, "", "Введение")
    paragraph(doc, "Цель работы - разработать систему управления базами данных на языке C++ с файловым хранением данных, SQL-подобным языком запросов и индексированием уникальных полей при помощи структуры B+-tree.")
    paragraph(doc, "Актуальность работы связана с изучением принципов системного программирования: организации файлового хранилища, обработки пользовательских команд, проверки корректности данных и применения индексных структур для ускорения поиска.")

    heading(doc, "1", "Постановка задачи")
    paragraph(doc, "Система должна поддерживать уровень системы, уровень базы данных и уровень таблицы. Пользователь взаимодействует с приложением через команды, завершающиеся символом точки с запятой. В пакетном режиме команды считываются из файла.")
    paragraph(doc, "Таблицы должны поддерживать типы int и string, значение NULL, модификаторы NOT_NULL, INDEXED и DEFAULT. Для SELECT результат должен выводиться в виде JSON-массива объектов.")
    paragraph(doc, "При выполнении условий выборки необходимо использовать имеющиеся индексы для оптимизации поиска. Для строковых значений требуется выполнять лексикографическое сравнение и поддерживать регулярные выражения LIKE.")

    heading(doc, "2", "Описание архитектуры")
    paragraph(doc, "Проект реализован как консольное приложение cwdb. Каталог data содержит базы данных. Внутри базы каждая таблица представлена отдельным каталогом, где размещаются файлы schema.txt, rows.tsv и history.log.")
    paragraph(doc, "Файл schema.txt хранит описание столбцов, их типы и ограничения. Файл rows.tsv хранит строки таблицы. Файл history.log используется для отката состояния таблицы командой REVERT.")
    add_table(doc)

    heading(doc, "3", "Индексирование B+-tree")
    paragraph(doc, "Индекс B+-tree реализован шаблонным классом BPlusTree. Внутренние узлы содержат разделяющие ключи, листовые узлы содержат пары ключ-номер строки. Листья связаны указателями next, что соответствует классической организации B+-tree.")
    paragraph(doc, "Для столбца INDEXED индекс строится автоматически. Ключом является значение столбца, значением является номер строки в файле таблицы. Такой подход выполняет требование не дублировать записи в индексной структуре.")
    paragraph(doc, "При условии вида column == value, где column является индексированным столбцом, поиск выполняется через B+-tree. При остальных условиях применяется последовательная проверка строк, потому что выражение может включать сложные предикаты, LIKE, BETWEEN, AND, OR и скобки.")

    heading(doc, "4", "Парсер и выполнение запросов")
    paragraph(doc, "Лексический анализатор выделяет имена, строковые литералы, числа, знаки операций и скобки. Ключевые слова обрабатываются регистронезависимо. Имена баз данных, таблиц и столбцов проверяются на соответствие требуемому формату.")
    paragraph(doc, "Реализованы команды CREATE DATABASE, DROP DATABASE, USE, CREATE TABLE, DROP TABLE, INSERT, UPDATE, DELETE, SELECT и REVERT. Для SELECT поддерживаются выбор всех столбцов, выбор перечисленных столбцов, алиасы AS и агрегатные функции SUM, COUNT, AVG.")
    paragraph(doc, "Условия WHERE вычисляются рекурсивным разбором выражений. Приоритет AND выше приоритета OR, а скобки позволяют явно задать порядок вычисления.")

    heading(doc, "5", "Надежность и журналирование")
    paragraph(doc, "Все данные сохраняются в файловой системе. После изменения строк таблицы данные записываются в rows.tsv, а индексы при загрузке и изменениях пересобираются по фактическим строкам.")
    paragraph(doc, "Перед UPDATE и DELETE состояние строк записывается в history.log с временной меткой. REVERT выбирает последнее состояние, созданное не позже указанного времени, и восстанавливает строки таблицы без копирования всей базы данных.")
    paragraph(doc, "Файл data/access.log содержит сведения о каждом запросе: время, идентификатор клиента local, обработчик cli, статус выполнения и тело запроса.")

    heading(doc, "6", "Тестирование")
    paragraph(doc, "Для проверки работы подготовлен пакетный сценарий examples/variant2_demo.sql. Он создает базу данных, таблицу с индексированным столбцом, вставляет строки, выполняет выборки, агрегаты, обновление и удаление.")
    demo = (ROOT / "examples" / "variant2_demo.sql").read_text(encoding="utf-8")
    add_listing(doc, "Листинг 1. Демонстрационный пакетный сценарий", demo)
    paragraph(doc, "Контрольная сборка выполнена командами cmake -G Ninja -S . -B build-ninja и cmake --build build-ninja. Демонстрационный сценарий выполняется командой build-ninja\\cwdb.exe examples\\variant2_demo.sql.")

    heading(doc, "7", "Руководство пользователя")
    paragraph(doc, "Для интерактивного режима необходимо запустить cwdb без аргументов. Для пакетного режима необходимо передать путь к файлу сценария первым аргументом командной строки.")
    paragraph(doc, "Каждая команда должна завершаться точкой с запятой. Результаты успешных SELECT-запросов выводятся в формате JSON, остальные успешные операции возвращают OK. Ошибки возвращаются строкой, начинающейся с ERROR.")

    heading(doc, "", "Вывод")
    paragraph(doc, "В результате работы реализована файловая СУБД для варианта 2. Программа поддерживает типизированные таблицы, ограничения целостности, уникальные B+-tree индексы, SQL-подобные команды, составные условия WHERE, агрегатные функции, журнал запросов и механизм отката таблицы.")

    heading(doc, "", "Список использованных источников")
    sources = [
        "1. ISO/IEC 14882:2017. Programming languages - C++ : International Standard. Geneva : International Organization for Standardization, 2017. 1605 p.",
        "2. Comer D. The Ubiquitous B-Tree // ACM Computing Surveys. 1979. Vol. 11, no. 2. P. 121-137.",
        "3. ГОСТ Р 7.0.100-2018. Библиографическая запись. Библиографическое описание. Общие требования и правила составления : национальный стандарт Российской Федерации : дата введения 2019-07-01. Москва : Стандартинформ, 2018. 124 с.",
        "4. ГОСТ 7.32-2017. Отчет о научно-исследовательской работе. Структура и правила оформления : межгосударственный стандарт : дата введения 2018-07-01. Москва : Стандартинформ, 2018. 32 с.",
        "5. Кормен Т. Х. Алгоритмы: построение и анализ / Т. Х. Кормен, Ч. И. Лейзерсон, Р. Л. Ривест, К. Штайн. 3-е изд. Москва : Вильямс, 2013. 1328 с.",
    ]
    for src in sources:
        paragraph(doc, src, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=False)

    heading(doc, "", "Приложение А")
    paragraph(doc, "Структура проекта:", align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=False)
    add_listing(doc, "Листинг 2. Основные файлы проекта", "\n".join([
        "CMakeLists.txt",
        "src/main.cpp",
        "examples/variant2_demo.sql",
        "tools/make_report.py",
        "README.md",
    ]))

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
